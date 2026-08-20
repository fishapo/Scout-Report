from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PROJECT_DOCUMENTATION.md"
OUTPUT = ROOT / "Scout-Report-Project-Documentation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
CONTENT_WIDTH_DXA = 9360


def set_font(run, name="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_cover(doc):
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("PROJECT DOCUMENTATION")
    set_font(run, size=12, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Scout Report")
    set_font(run, size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("Agricultural pest and disease reporting platform")
    set_font(run, size=15, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Code-verified functionality, current gaps, and recommended roadmap")
    set_font(run, size=11, italic=True, color=MUTED)
    for _ in range(10):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 2.0.0 | Audit date: 6 August 2026")
    set_font(run, size=10, color=MUTED)
    doc.add_page_break()


def add_inline_markdown(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9.5, color=NAVY)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    headers = rows[0]
    body = [row for row in rows[2:] if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
    return headers, body, i


def add_table(doc, headers, rows):
    columns = len(headers)
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    widths = [CONTENT_WIDTH_DXA // columns] * columns
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for cell, value in zip(header_cells, headers):
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_markdown(p, value)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown_body(doc, markdown):
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
                for idx, code_line in enumerate(code_lines):
                    run = p.add_run(code_line + ("\n" if idx < len(code_lines) - 1 else ""))
                    set_font(run, name="Consolas", size=9.5, color=NAVY)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            headers, rows, i = parse_table(lines, i)
            add_table(doc, headers, rows)
            continue
        if stripped.startswith("# "):
            # Source title belongs to the cover already.
            i += 1
            continue
        match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if match:
            level = len(match.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline_markdown(p, match.group(2))
            i += 1
            continue
        if stripped.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_geometry(table, [CONTENT_WIDTH_DXA])
            cell = table.cell(0, 0)
            set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, stripped[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or numbered:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline_markdown(p, (bullet or numbered).group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)
        i += 1


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Scout Report | Project Documentation")
    set_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    add_cover(doc)
    add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "Scout Report Project Documentation"
    doc.core_properties.subject = "Functionality, pending work, and future roadmap"
    doc.core_properties.author = "Scout Report Project"
    doc.core_properties.comments = "Generated from the code-verified project documentation."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
